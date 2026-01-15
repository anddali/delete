"""
File parsing service for extracting text from various document formats.

Supports: PDF, DOCX, TXT, MD, HTML, JSON, CSV, XML
"""

import csv
import io
import json
import os
from typing import Optional
from xml.etree import ElementTree

import structlog

logger = structlog.get_logger()


class FileParser:
    """Base file parser interface."""
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        """Check if this parser supports the given extension."""
        raise NotImplementedError
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse file content and return extracted text."""
        raise NotImplementedError


class TextParser(FileParser):
    """Parser for plain text files (.txt, .md)."""
    
    SUPPORTED_EXTENSIONS = {".txt", ".md"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in TextParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse text/markdown files."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode text file: {filename}")


class HTMLParser(FileParser):
    """Parser for HTML files."""
    
    SUPPORTED_EXTENSIONS = {".html", ".htm"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in HTMLParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse HTML and extract text content."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback to basic text extraction
            text = content.decode("utf-8", errors="ignore")
            # Remove script and style tags (basic)
            import re
            text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<[^>]+>", " ", text)
            text = re.sub(r"\s+", " ", text)
            return text.strip()
        
        soup = BeautifulSoup(content, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        # Get text
        text = soup.get_text(separator="\n")
        
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines()]
        text = "\n".join(line for line in lines if line)
        
        return text


class PDFParser(FileParser):
    """Parser for PDF files using pdfplumber."""
    
    SUPPORTED_EXTENSIONS = {".pdf"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in PDFParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse PDF and extract text content."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF parsing. "
                "Install with: pip install pdfplumber"
            )
        
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(
                        "Failed to extract text from PDF page",
                        page=page_num,
                        error=str(e),
                        filename=filename,
                    )
        
        if not text_parts:
            raise ValueError(f"Could not extract any text from PDF: {filename}")
        
        return "\n\n".join(text_parts)


class DOCXParser(FileParser):
    """Parser for Microsoft Word documents (.docx)."""
    
    SUPPORTED_EXTENSIONS = {".docx"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in DOCXParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse DOCX and extract text content."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )
        
        doc = Document(io.BytesIO(content))
        
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        if not text_parts:
            raise ValueError(f"Could not extract any text from DOCX: {filename}")
        
        return "\n\n".join(text_parts)


class JSONParser(FileParser):
    """Parser for JSON files."""
    
    SUPPORTED_EXTENSIONS = {".json"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in JSONParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse JSON and convert to readable text."""
        text = content.decode("utf-8")
        data = json.loads(text)
        
        def extract_text(obj, depth=0) -> list:
            """Recursively extract text from JSON structure."""
            texts = []
            indent = "  " * depth
            
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if isinstance(value, (str, int, float, bool)):
                        texts.append(f"{indent}{key}: {value}")
                    else:
                        texts.append(f"{indent}{key}:")
                        texts.extend(extract_text(value, depth + 1))
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    if isinstance(item, (str, int, float, bool)):
                        texts.append(f"{indent}- {item}")
                    else:
                        texts.append(f"{indent}[{i}]:")
                        texts.extend(extract_text(item, depth + 1))
            else:
                texts.append(f"{indent}{obj}")
            
            return texts
        
        return "\n".join(extract_text(data))


class CSVParser(FileParser):
    """Parser for CSV files."""
    
    SUPPORTED_EXTENSIONS = {".csv"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in CSVParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse CSV and convert to readable text."""
        text = content.decode("utf-8")
        reader = csv.reader(io.StringIO(text))
        
        rows = list(reader)
        if not rows:
            return ""
        
        # Use first row as headers
        headers = rows[0] if rows else []
        
        text_parts = []
        for i, row in enumerate(rows[1:], start=1):
            if headers:
                row_text = ", ".join(
                    f"{headers[j]}: {cell}" 
                    for j, cell in enumerate(row) 
                    if j < len(headers) and cell.strip()
                )
            else:
                row_text = ", ".join(cell for cell in row if cell.strip())
            
            if row_text:
                text_parts.append(f"Row {i}: {row_text}")
        
        return "\n".join(text_parts)


class XMLParser(FileParser):
    """Parser for XML files."""
    
    SUPPORTED_EXTENSIONS = {".xml"}
    
    @staticmethod
    def can_parse(extension: str) -> bool:
        return extension.lower() in XMLParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse(content: bytes, filename: str) -> str:
        """Parse XML and extract text content."""
        root = ElementTree.fromstring(content)
        
        def extract_text(element, depth=0) -> list:
            """Recursively extract text from XML elements."""
            texts = []
            indent = "  " * depth
            
            # Get element text
            if element.text and element.text.strip():
                texts.append(f"{indent}{element.tag}: {element.text.strip()}")
            elif element.attrib:
                attrs = ", ".join(f"{k}={v}" for k, v in element.attrib.items())
                texts.append(f"{indent}{element.tag} ({attrs})")
            
            # Recurse into children
            for child in element:
                texts.extend(extract_text(child, depth + 1))
            
            return texts
        
        return "\n".join(extract_text(root))


# Registry of all parsers
PARSERS = [
    TextParser,
    HTMLParser,
    PDFParser,
    DOCXParser,
    JSONParser,
    CSVParser,
    XMLParser,
]


def get_parser(extension: str) -> Optional[FileParser]:
    """Get the appropriate parser for a file extension."""
    for parser in PARSERS:
        if parser.can_parse(extension):
            return parser
    return None


def parse_file(content: bytes, filename: str) -> str:
    """
    Parse a file and extract its text content.
    
    Args:
        content: Raw file bytes
        filename: Original filename (used to determine type)
    
    Returns:
        Extracted text content
    
    Raises:
        ValueError: If file type is not supported or parsing fails
    """
    ext = os.path.splitext(filename)[1].lower()
    
    parser = get_parser(ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {ext}")
    
    logger.info("Parsing file", filename=filename, parser=parser.__name__)
    
    try:
        text = parser.parse(content, filename)
        logger.info(
            "File parsed successfully",
            filename=filename,
            text_length=len(text),
        )
        return text
    except Exception as e:
        logger.error("Failed to parse file", filename=filename, error=str(e))
        raise


def get_supported_extensions() -> set:
    """Get all supported file extensions."""
    extensions = set()
    for parser in PARSERS:
        if hasattr(parser, "SUPPORTED_EXTENSIONS"):
            extensions.update(parser.SUPPORTED_EXTENSIONS)
    return extensions
